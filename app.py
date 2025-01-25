from flask import Flask, render_template, request, redirect, url_for, send_from_directory
import os
import fitz  # PyMuPDF
from PIL import Image
import google.generativeai as genai
from docx import Document
import io
import re
import PyPDF2
import cv2

genai.configure(api_key="insira key")
model = genai.GenerativeModel("gemini-1.5-flash")

app = Flask(__name__, static_folder='static')

UPLOAD_FOLDER = 'uploads'
RESULT_FOLDER = 'results'
IMAGES_FOLDER = 'images'

# pastas
for folder in [UPLOAD_FOLDER, RESULT_FOLDER, IMAGES_FOLDER]:
    os.makedirs(folder, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULT_FOLDER'] = RESULT_FOLDER
app.config['IMAGES_FOLDER'] = IMAGES_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/img_to_doc')
def img_to_doc():
    return render_template('img_to_doc.html')

@app.route('/pdf_to_doc')
def pdf_to_doc():
    return render_template('pdf_to_doc.html')

#funções para tratamento do texto
def clean_text(text):
    return re.sub(r'[\x00-\x1F\x7F]', '', text)

def fix_line_breaks(text):
    return re.sub(r'(?<!\n)\n(?!\n)', ' ', text)

# Função para extrair texto em blocos do PDF
def extract_text_with_blocks(filepath):
    pdf_document = fitz.open(filepath)
    full_text = ""
    for page in pdf_document:
        blocks = page.get_text("blocks") 
        if blocks:
            for block in blocks:
                block_text = block[4].strip()
                full_text += block_text + "\n"
        full_text += "\n\n" 
    pdf_document.close()
    return clean_text(full_text)

# tirar imagens do pdf***
def extract_images_from_pdf(filepath, output_folder):
    pdf_document = fitz.open(filepath)
    image_paths = []
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image = Image.open(io.BytesIO(image_bytes))
            output_path = os.path.join(output_folder, f"page_{page_num + 1}_img_{img_index + 1}.{image_ext}")
            image.save(output_path)
            image_paths.append(output_path)
    pdf_document.close()
    return image_paths

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return "Nenhum arquivo enviado", 400

    file = request.files['file']
    if file.filename == '':
        return "Nenhum arquivo selecionado", 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        block_text = extract_text_with_blocks(filepath)
        images = extract_images_from_pdf(filepath, app.config['IMAGES_FOLDER'])

        image_text = ""
        for image_path in images:
            with open(image_path, "rb") as img_file:
                img = Image.open(img_file)
                response = model.generate_content([
                    "Transcreva apenas o texto presente nesta imagem:",
                    img
                ])
                image_text += response.text.strip() + "\n\n"

        # combinação do que é extraido 
        combined_text = f"{block_text}\n\n---\n\n{image_text}"
        combined_text = fix_line_breaks(combined_text)

        #  Google GenerativeAI para reformatar
        response = model.generate_content([
            "Reestruture o seguinte texto extraído de um PDF, mantendo as quebras de linha e a estrutura original:",
            combined_text
        ])

        # resultado
        doc = Document()
        doc.add_paragraph(response.text.strip())
        doc_path = os.path.join(app.config['RESULT_FOLDER'], "converted_document.docx")
        doc.save(doc_path)

        return redirect(url_for('download_file', filename="converted_document.docx"))

    except Exception as e:
        return f"Erro ao processar o PDF: {e}", 500

@app.route('/upload_image', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return "Nenhum arquivo enviado", 400

    file = request.files['file']
    if file.filename == '':
        return "Nenhum arquivo selecionado", 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        image = cv2.imread(filepath)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 102, 255, cv2.THRESH_BINARY_INV)
        processed_image_path = os.path.join(app.config['UPLOAD_FOLDER'], "imagem_processada.png")
        cv2.imwrite(processed_image_path, thresh)

        processed_image = Image.open(processed_image_path)
        response = model.generate_content([
            "O que está escrito? Apenas transcreva o texto sem comentar nada",
            processed_image
        ])

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        text_doc = response.text
        doc.add_paragraph(text_doc)

        doc_path = os.path.join(app.config['RESULT_FOLDER'], "document.docx")
        doc.save(doc_path)

        return redirect(url_for('download_file', filename="document.docx"))

    except Exception as e:
        return f"Erro ao processar a imagem: {e}", 500

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['RESULT_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)